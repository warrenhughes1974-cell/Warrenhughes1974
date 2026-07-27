"""Builder: CSO PUA conversion briefing in plain English."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

OUT_PATH = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "Issue_Log_Items",
    "PUA_CSO_Conversion_Briefing.docx",
)


def set_cell_shading(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>')
    tcPr.append(shd)


def add_para(doc, text, bold=False, italic=False, size=11, space_after=8, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(11)


def style_table(table, header=True):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(2)
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(9)
                    if i == 0 and header:
                        run.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            if i == 0 and header:
                set_cell_shading(cell, "1F3A5F")
            elif i % 2 == 1:
                set_cell_shading(cell, "F2F5F8")


def fill_table(table, rows):
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            table.rows[i].cells[j].text = val
    style_table(table)


def main():
    os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(8)
    style.paragraph_format.line_spacing = 1.15

    for i in range(1, 4):
        hs = doc.styles[f"Heading {i}"]
        hs.font.name = "Calibri"
        hs.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
        hs.font.size = Pt(16 if i == 1 else 13 if i == 2 else 11)

    # Cover
    for _ in range(2):
        doc.add_paragraph()

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Paid-Up Additions (PUA)")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    r.font.name = "Calibri"

    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = st.add_run(
        "How They Convert from LifePRO\nand How They Work in QLAdmin"
    )
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(0x3D, 0x5A, 0x80)
    r.font.name = "Calibri"

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run("CSO Conversion Briefing\nPrepared: July 27, 2026")
    r.font.size = Pt(11)
    r.font.name = "Calibri"
    r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_page_break()

    # 1
    doc.add_heading("1. Purpose", level=1)
    add_para(
        doc,
        "This briefing explains how Paid-Up Additions are converted from LifePRO into "
        "QLAdmin for the CSO book, and how QLAdmin calculates their values under the "
        "current conversion design.",
    )

    # 2
    doc.add_heading("2. What a Paid-Up Addition is", level=1)
    add_para(
        doc,
        "A Paid-Up Addition is extra paid-up life insurance bought on a participating "
        "permanent policy, usually with dividends. It increases the death benefit and "
        "adds to cash value. It is not a new policy the customer bought on its own.",
    )
    add_para(doc, "In LifePRO", bold=True, space_after=4)
    add_bullets(
        doc,
        [
            "PUAs show up as their own coverage rows on the policy (for example 960 PO PUA or 670 PUA).",
            "They carry a unit amount, often with decimal places, and may have their own cash-value tables in LifePRO.",
            "They sit next to the base coverage on the same policy.",
        ],
    )
    add_para(doc, "In QLAdmin", bold=True, space_after=4)
    add_bullets(
        doc,
        [
            "PUAs are loaded as an extra coverage on the same policy, not as a separate policy.",
            'The coverage plan code is built from the first four characters of the base plan, plus "PA". Example: base plan 1960PO becomes PUA plan 1960PA.',
            "We do not create a separate PA plan in the plan file. When that PA plan is missing, QLAdmin uses the base plan's rates to calculate the PUA values — as long as the PUA coverage fields are set correctly.",
            "On an active base coverage, the PUA status is Paid Up (41). If the base is on Extended Term or Reduced Paid-Up, the PUA is terminated (54).",
        ],
    )

    # 3
    doc.add_heading("3. How QLAdmin calculates PUA values", level=1)
    add_para(
        doc,
        "QLAdmin treats a PUA as an extra coverage that is already paid up. Good results "
        "depend on two things: the PUA coverage fields must be set correctly, and the "
        "base plan must have usable interest and mortality assumptions.",
    )
    add_para(doc, "Which rates does QLAdmin use?", bold=True, space_after=4)
    add_bullets(
        doc,
        [
            "If a PA plan exists in the plan file with its own cash-value and reserve rates, QLAdmin uses those rates.",
            "If that PA plan is not in the plan file, QLAdmin uses the base plan's rates instead.",
            "Our conversion uses the second approach: the coverage is labeled with a PA code (for example 1960PA), but that code is not added to the plan file, so the base plan drives the calculation.",
        ],
    )
    add_para(
        doc,
        "That way, when interest or mortality changes on the base plan, the PUA follows "
        "along without a separate set of PA plans and rates to maintain.",
    )

    # 4
    doc.add_heading("4. Conversion rules we follow", level=1)
    table = doc.add_table(rows=9, cols=2)
    table.style = "Table Grid"
    fill_table(
        table,
        [
            ("Area", "Rule"),
            (
                "Plan file",
                "Do not add PA plans to the plan file. The PA code appears only on the policy coverage.",
            ),
            (
                "Factors",
                "Do not load special conversion factors for PUAs. Let QLAdmin calculate the values.",
            ),
            (
                "Status",
                "PUA status is Paid Up (41) when the base is active. If the base is Extended Term (44) or Reduced Paid-Up (45), the PUA is terminated (54).",
            ),
            (
                "Effective date and age",
                "PUA effective date and issue age match the base coverage.",
            ),
            (
                "Paid-up date",
                "PUA paid-up date equals the PUA effective date.",
            ),
            (
                "Duration",
                "Years since issue on the PUA is recalculated from that inherited effective date so it lines up with the base.",
            ),
            (
                "Expiry and plan code",
                'Expiry matches the base. Plan code is the first four characters of the base plan plus "PA".',
            ),
            (
                "Other coverages",
                "These rules apply only to Paid-Up Additions — not to ADB, waiver, term riders, or other coverages.",
            ),
        ],
    )
    add_para(doc, "")

    # 5
    doc.add_heading("5. Why the coverage fields matter", level=1)
    add_para(
        doc,
        "Even when QLAdmin uses the base plan rates, it will not calculate good PUA values "
        "if the PUA coverage looks like a late-issue active rider. Here is the before-and-after "
        "picture on sample policy 9010310404C (base 1960PO / PUA 1960PA):",
    )
    table = doc.add_table(rows=7, cols=3)
    table.style = "Table Grid"
    fill_table(
        table,
        [
            ("Coverage field", "Incorrect setup", "Correct conversion"),
            ("Status", "22 Active", "41 Paid Up"),
            ("Effective date", "01/28/2011 (when PUA was issued)", "01/28/1969 (same as base)"),
            ("Issue age", "68", "26 (same as base)"),
            ("Paid-up date", "01/28/2046", "01/28/1969 (same as effective date)"),
            ("Years since issue", "15", "57 (same as base)"),
            ("Plan code", "1960PA", "1960PA (naming pattern)"),
        ],
    )
    add_para(doc, "")
    add_para(
        doc,
        "After the data is loaded: run Data Admin, then rebuild cash values on the policy. "
        "Correct dollar results also require that the base plan has non-zero cash-value "
        "and reserve interest set up.",
        italic=True,
    )

    # 6
    doc.add_heading("6. How the conversion works", level=1)

    doc.add_heading("6.1 Step by step", level=2)
    add_bullets(
        doc,
        [
            "1. Find LifePRO coverage rows that are Paid-Up Additions.",
            "2. Recognize them through the product crosswalk so the converter knows they are PUAs.",
            "3. Convert the base coverage first and remember its key fields for that policy.",
            '4. Convert the PUA coverage: set the plan code to the base plan\'s first four characters plus "PA", copy effective date / age / expiry from the base, set the paid-up date to the effective date, and set status to Paid Up (or terminated if the base is on nonforfeiture).',
            "5. Recalculate years since issue from the inherited effective date.",
            "6. Do not add that PA plan code to the plan file.",
            "7. After load in QLAdmin, run Data Admin and rebuild cash values so the system calculates PUA values from the base plan.",
            "8. Dividend history that bought PUAs is also converted (LifePRO dividend-to-PUA transactions become benefit-history type 4).",
        ],
    )

    doc.add_heading("6.2 Which products are treated as PUAs", level=2)
    add_para(
        doc,
        "The converter treats these product codes as Paid-Up Additions:",
        space_after=4,
    )
    add_para(
        doc,
        "280PUA, 121PUA, 1970PA, 170PUA, 165PUA, 185PUA, 261PUA, 1OLPUA, 1POPUA, 265PUA, "
        "and LifePRO label 970 PUA.",
        italic=True,
    )
    add_para(
        doc,
        "Only these products get the PUA inheritance rules. Other riders keep their own "
        "issue dates, ages, and statuses.",
    )

    doc.add_heading("6.3 Fields set on the PUA coverage", level=2)
    table = doc.add_table(rows=8, cols=2)
    table.style = "Table Grid"
    fill_table(
        table,
        [
            ("Field", "What we set"),
            (
                "Plan code",
                'First four characters of the base plan + "PA" (example: 1960PO → 1960PA)',
            ),
            ("Effective date", "Same as the base coverage"),
            ("Issue age", "Same as the base coverage"),
            ("Expiry date", "Same as the base coverage"),
            ("Paid-up date", "Same as the PUA effective date"),
            (
                "Status",
                "41 Paid Up if the base is active; 54 terminated if the base is Extended Term or Reduced Paid-Up",
            ),
            (
                "Years since issue",
                "Recalculated from the inherited effective date so it matches the base",
            ),
        ],
    )
    add_para(doc, "")

    doc.add_heading("6.4 Processing order", level=2)
    add_para(
        doc,
        "A PUA row in the LifePRO extract may appear before or after the base coverage. "
        "The converter remembers the base fields as it goes, holds PUA rows until the base "
        "for that policy is available, then applies the inheritance rules. That way every "
        "PUA is aligned to its own base coverage.",
    )

    # 7
    doc.add_heading("7. Face amount, participation, and dividends", level=1)

    doc.add_heading("7.1 Face amount", level=2)
    add_bullets(
        doc,
        [
            "LifePRO units become the coverage units in QLAdmin. The value per unit is 1,000.",
            "Face amount = units × 1,000. Example on policy 9010448806C: 5.75296 units × 1,000 = $5,752.96.",
            "The conversion keeps five decimal places on the units. The Coverage screen may round the amount to whole dollars for display, but calculations and payments use the stored units.",
        ],
    )

    doc.add_heading("7.2 Participating / non-participating", level=2)
    add_para(
        doc,
        "A Paid-Up Addition is not a participating coverage. When QLAdmin adds a PA "
        "coverage, it sets the participating flag (PAR / MPAR) to 0 on that coverage, "
        "even if the base plan is participating. The base coverage can still be "
        "participating; the PUA coverage itself is not.",
    )

    doc.add_heading("7.3 Dividend purchases", level=2)
    add_para(
        doc,
        "When a policy earns a dividend and the owner uses that dividend to buy more "
        "paid-up insurance, LifePRO records the purchase. We convert those records into "
        "QLAdmin benefit history (type 4) so you can see each year's dividend that was "
        "used to buy PUAs. That history supports the PUA coverage amount on the policy; "
        "it is not the same as dividend money left on deposit, which stays as a deposit "
        "balance and is not paid-up insurance.",
    )

    # 8
    doc.add_heading("8. What happens on Extended Term or Reduced Paid-Up", level=1)
    add_para(
        doc,
        "When the base coverage is on Extended Term or Reduced Paid-Up, other coverages — "
        "including Paid-Up Additions — should not stay in force as Paid Up.",
    )
    add_bullets(
        doc,
        [
            "If the base status is Extended Term (44) or Reduced Paid-Up (45), the PUA status is set to terminated (54).",
            "That keeps PUAs from remaining Paid Up on nonforfeiture policies.",
        ],
    )

    # 9
    doc.add_heading("9. How many PUAs are in the book", level=1)
    add_para(
        doc,
        "There are about 10 LifePRO PUA products and about 495 PUA coverage rows in the "
        "CSO conversion. The main groups look like this:",
    )
    table = doc.add_table(rows=7, cols=4)
    table.style = "Table Grid"
    fill_table(
        table,
        [
            ("QLAdmin coverage plan", "LifePRO PUA source", "Approx. policies", "In plan file?"),
            ("1708PA / 1705PA", "670 PUA", "~415", "No (by design)"),
            ("1960PA", "960 PO / OL / 65 / LP PUA", "~71", "No (by design)"),
            ("280EPA", "980 PUA", "3", "No"),
            ("221EPA", "621 PUA", "1", "No"),
            ("2665PA", "665 PUA", "1", "No"),
            ("1970PA / 261PUA", "970 PUA / 961 PUA", "1 each", "Still uses inheritance from the base"),
        ],
    )
    add_para(doc, "")
    add_para(
        doc,
        "One PA coverage code can cover more than one LifePRO PUA source when those "
        "sources sit under the same base plan family. That works under the current design "
        "because QLAdmin calculates from the base plan rather than from separate PA rate tables.",
        italic=True,
    )

    # 10
    doc.add_heading("10. How we check the results", level=1)
    add_bullets(
        doc,
        [
            "We check that PUA status, dates, age, paid-up date, and years since issue follow the rules above — including terminated status on Extended Term and Reduced Paid-Up.",
            "We check that unit amounts keep five decimal places.",
            "We check that each PUA coverage has participating flag (PAR / MPAR) set to 0.",
            "We confirm synthetic PA plan codes such as 1960PA are not present in the plan file.",
            "Plan-file rules also prevent reserved PA-style endings from being used as regular plan codes.",
        ],
    )
    add_para(
        doc,
        "Some batch reports may still note that PA coverage codes are not in the plan file. "
        "Under this design that is expected and does not block the load.",
    )

    # 11
    doc.add_heading("11. Sample policies", level=1)
    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    fill_table(
        table,
        [
            ("Policy", "Why it matters", "What to look at"),
            (
                "9010310404C",
                "Main sample — base 1960PO, PUA 1960PA",
                "Coverage fields; Data Admin and rebuild cash values; base plan interest setup",
            ),
            (
                "9010448806C",
                "Unit precision — PUA face $5,752.96",
                "Stored units 5.75296 versus whole-dollar display on the Coverage screen",
            ),
            (
                "Policies on Extended Term or Reduced Paid-Up with a PUA",
                "PUA should be terminated (54), not Paid Up (41)",
                "Confirm the PUA is not left in force on a nonforfeiture policy",
            ),
        ],
    )
    add_para(doc, "")

    add_para(
        doc,
        "Document status: process briefing for the CSO Paid-Up Addition conversion as of July 2026.",
        italic=True,
        size=10,
    )

    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    main()
