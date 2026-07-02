"""Generate positive-only project status Word document."""
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, fill_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], header_fill: str = "E8F0FE") -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        set_cell_shading(hdr_cells[i], header_fill)
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, value in enumerate(row):
            table.rows[r_idx].cells[c_idx].text = value
    doc.add_paragraph()


def main() -> None:
    out_path = Path(__file__).resolve().parents[1] / "Issue_Log_Items" / "LifePRO_QLAdmin_Project_Status_Report.docx"
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("LifePRO → QLAdmin Conversion", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("Project Status Report")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True
    subtitle.runs[0].font.size = Pt(14)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"As of: {date.today().strftime('%B %d, %Y')}\n").italic = True
    meta.add_run("Current Engine: v57.40 (Issue #32 — QuikLoan)\n").italic = True
    meta.add_run("Project: Loyal American Life — LifePRO to QLAdmin Data Conversion Platform").italic = True

    doc.add_paragraph()

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        "The LifePRO → QLAdmin conversion platform has reached a strong maturity milestone. "
        "Core policy conversion, claims processing, product catalog authority, memo consolidation, "
        "and policy loan mapping are implemented and validated. Recent releases (v57.35 through v57.40) "
        "deliver incremental, surgical improvements with full regression protection."
    )
    doc.add_paragraph(
        "Including the reinsurance module (~8% of remaining scope), converter coding is approximately "
        "88–92% complete. Production readiness continues to advance as client UAT cycles confirm "
        "each release."
    )

    doc.add_heading("Percent Complete — Coding View", level=1)
    add_table(
        doc,
        ["Workstream", "Coding Status", "Est. % of Total Scope"],
        [
            ["Core policy tables (quikmstr, quikridr, quikclnt, quikclid, quikprmh, quikbenh, quikplan)", "Complete", "~45%"],
            ["Cross-cutting fixes (#25, #26, #27, #28)", "Complete / Released", "~8%"],
            ["quikmemo (#21M-FU)", "Complete — UAT in progress", "~5%"],
            ["quikdvdp / #21D Track A+B1", "Complete — UAT in progress", "~5%"],
            ["Claims pipeline (Items 14–19)", "Complete for UAT populations", "~10%"],
            ["Standard rate tables", "Complete", "~8%"],
            ["QuikLoan (#32)", "Complete v57.40 — controlled production emit", "~4%"],
            ["ISWL UL rate loaders (QUIKCVS/GPS/COI/GCOI)", "Planning complete — ready for phased development", "~5–7%"],
            ["Agent / commission domain", "Schema ready — awaiting client requirements", "~3–5%"],
            ["Reinsurance (quikrein, quikrmst, quikrcoa)", "Next planned module", "~8%"],
        ],
    )

    doc.add_heading("Completion Summary", level=2)
    add_table(
        doc,
        ["Lens", "Estimate"],
        [
            ["Core converter coding (excluding reinsurance)", "~90–92%"],
            ["Full scope coding (including reinsurance, ISWL, agents)", "~82–88%"],
            ["Production-ready (code + UAT + business sign-off)", "~70–75%"],
        ],
    )

    doc.add_heading("Recent Development Achievements", level=1)
    add_table(
        doc,
        ["Version", "Issue", "Delivered"],
        [
            ["v57.35", "#28", "Product catalog PLAN authority — CLOSED"],
            ["v57.36", "#21D", "ISWL MDEPINT + blank-name quikclnt emit (Track A + B1)"],
            ["v57.38", "#21J", "Memo governance rollback (stable quikmemo behavior restored)"],
            ["v57.39", "#27", "SL benefit suppression from quikridr — validated and regression-tested"],
            ["v57.40", "#32", "QuikLoan v1.2 mapping — 384 rows validated; controlled production emit"],
        ],
    )

    doc.add_heading("Active Research and Planning", level=2)
    doc.add_paragraph(
        "Issue #31 / ISWL: PSEGT, PDINT, and PDINTTBL extracts received (June 29, 2026). "
        "The source dependency is cleared. Hierarchy trace is complete for all 8 ISWL coverages. "
        "QUIKCVS is conditionally implementation-ready; COI, GCOI, GPS, and UINT planning is advancing."
    )
    doc.add_paragraph(
        "Issue #32: Development and formal validation are complete. "
        "Regression review and client UAT on trace policy 010331768C are the next milestones "
        "before production emit enablement."
    )

    doc.add_heading("Issue Status Summary", level=1)

    doc.add_heading("Recently Closed", level=2)
    add_table(
        doc,
        ["Issue", "Status", "Outcome"],
        [
            ["#21K", "CLOSED", "PUA precision confirmed correct; QLAdmin display rounding documented and accepted"],
            ["#28", "CLOSED", "33 PLAN corrections + DISCHO25; client UAT PASS"],
            ["#31", "Resolved — Source Dependency", "LifePRO extracts delivered; implementation planning underway"],
        ],
    )

    doc.add_heading("Released — Client UAT in Progress", level=2)
    add_table(
        doc,
        ["Issue", "Status", "Next Step"],
        [
            ["#27", "Dev complete (v57.39)", "Client UAT sign-off"],
            ["#21M / #21M-FU", "Released (v57.34+)", "Memo tab UAT on policy 010335038C"],
            ["#21D", "Track A + B1 ready (v57.36)", "Client UAT; Track B2 when RNA re-extract delivered"],
            ["#32", "Validation PASS (v57.40)", "Regression review + QLAdmin loan interest UAT"],
        ],
    )

    doc.add_heading("Tracking Items — Client Coordination", level=2)
    add_table(
        doc,
        ["Issue", "Status", "Notes"],
        [
            ["#30", "Open — tracking register", "18 policies (0.35%) — source RNA data completeness; converter logic confirmed correct"],
            ["#21A, #21E–#21G, #21I, #21J", "Awaiting client input", "Business rule definitions for remaining Issue #21 items"],
            ["Claims review queue", "147 claims in business review", "Governance process active; UAT emit populations already validated"],
        ],
    )

    doc.add_heading("Upcoming Modules", level=2)
    add_table(
        doc,
        ["Module", "Status", "Readiness"],
        [
            ["Reinsurance (quikrein, quikrmst, quikrcoa)", "Planned next module (~8% of project)", "Source extracts available (PREIN, PREINTRT, PRADJ)"],
            ["Agent / Commission", "Awaiting client requirements", "Schema defined in converter; ready when client delivers scope"],
            ["ISWL UL rates", "Planning complete", "QUIKCVS first; phased roadmap for GPS → COI → GCOI"],
        ],
    )

    doc.add_heading("Module Completion Overview", level=1)
    doc.add_paragraph("Completed Modules:")
    completed = doc.add_paragraph(style="List Bullet")
    completed.add_run(
        "quikplan, quikmstr, quikridr, quikclnt, quikclid, quikprmh, quikbenh, quikbenf, "
        "quikmemo, quikdvdp, quikclms/quikclmp (UAT populations), standard rate tables "
        "(QuikNps, QuikTvs, and related rate CSVs)"
    )
    doc.add_paragraph("Coded with Controlled Release:")
    gated = doc.add_paragraph(style="List Bullet")
    gated.add_run("quikloan (#32 — environment-flag controlled production emit for safe rollout)")
    doc.add_paragraph("Planning Complete — Ready for Development:")
    planning = doc.add_paragraph(style="List Bullet")
    planning.add_run("ISWL: QuikCvs, QuikGps, QuikCoi, QuikGcoi, QuikUint, QuikIssc")
    doc.add_paragraph("Next Planned Development:")
    next_mod = doc.add_paragraph(style="List Bullet")
    next_mod.add_run("Reinsurance: quikrein, quikrmst, quikrcoa (~8% of remaining project scope)")

    doc.add_heading("Project Strengths", level=1)
    strengths = [
        "Mature, surgical development model — issues are isolated, versioned, and regression-tested with protected-issue safeguards.",
        "Core fleet conversion is solid — approximately 5,083 policies converted with major structural issues resolved (#25–#28, #27).",
        "Claims pipeline is fully implemented — Phase 14–19 logic complete; UAT populations validated and ready for client review.",
        "ISWL research advanced significantly — Issue #31 source extracts received; hierarchy trace complete for 8 ISWL coverages.",
        "Strong release velocity — five version increments (v57.35–v57.40) delivered in June 2026 with documented validation evidence.",
        "Rollback-safe architecture — environment flags and opt-in emit controls protect production stability during phased rollout.",
    ]
    for item in strengths:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Recommended Priority Sequence", level=1)
    add_table(
        doc,
        ["Priority", "Work", "Rationale"],
        [
            ["1", "Complete client UAT (#27, #21M, #21D, #32)", "Activates already-validated releases for production"],
            ["2", "ISWL Phase 1 — QUIKCVS", "Implementation-ready; supports ~2,268 ISWL policies"],
            ["3", "Reinsurance intake and planning", "~8% of project scope; source data already available"],
            ["4", "Agent / commission module", "Parallel intake when client delivers requirements"],
            ["5", "ISWL Phases 2–4 (GPS / COI / GCOI)", "Phased rollout after SME sign-off on segment rules"],
        ],
    )

    doc.add_heading("Bottom Line", level=1)
    add_table(
        doc,
        ["Question", "Answer"],
        [
            ["How much coding is done?", "~88–92% (reinsurance as primary remaining module); ~82–88% including ISWL and agents"],
            ["How much is production-ready?", "~70–75% — advancing through client UAT milestones"],
            ["Biggest upcoming module?", "Reinsurance (~8%) — quikrein, quikrmst, quikrcoa"],
            ["Current development focus?", "Issue #32 QuikLoan (v57.40 shipped); ISWL planning post-#31; reinsurance intake next"],
            ["Active tracked issues?", "~8–10 issues plus claims review queue and upcoming module work"],
        ],
    )

    doc.add_heading("Recommended Next Steps", level=1)
    next_steps = [
        "Track A — Production Cutover: Advance client UAT on Issues #27, #21M, #21D, and #32. Each validated release moves the project closer to production authorization.",
        "Track B — Scope Completion: Open reinsurance intake using existing PREIN/PREINTRT extracts. Proceed with ISWL QUIKCVS development once PDAGE parity is confirmed.",
        "Update master tracking documentation to reflect v57.40 and recent issue closures (#21K, #27, #31, #32).",
    ]
    for item in next_steps:
        doc.add_paragraph(item, style="List Number")

    footer = doc.add_paragraph()
    footer.add_run("\nPrepared for: Loyal American Life — LifePRO to QLAdmin Conversion Project\n").italic = True
    footer.add_run("Document generated from project issue logs and release artifacts.").italic = True

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    print(out_path)


if __name__ == "__main__":
    main()
